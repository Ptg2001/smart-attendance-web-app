from flask import Flask, render_template, request, redirect, url_for, session, send_file
from pymongo import MongoClient
import cv2
import face_recognition
import numpy as np
import speech_recognition as sr
import pandas as pd
import os
from io import BytesIO
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle
from reportlab.lib import colors
from docx import Document
from docx.shared import Pt
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
import base64
import io
from dotenv import load_dotenv

# Load environment variables from .env file
load_dotenv()

app = Flask(__name__)

# MongoDB Atlas configuration
MONGO_URI = os.getenv("MONGO_URI")
client = MongoClient(MONGO_URI)
db = client.attendance_db

# Mock user for simplicity
users = {'admin': 'password'}

def get_known_faces():
    face_encodings = []
    face_names = []
    records = db.face_encodings.find()
    for record in records:
        face_encodings.append(np.array(record["encoding"]))
        face_names.append(record["name"])
    return face_encodings, face_names

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        username = request.form['username']
        password = request.form['password']
        if username in users and users[username] == password:
            session['username'] = username
            return redirect(url_for('dashboard'))
    return render_template('login.html')

@app.route('/dashboard')
def dashboard():
    if 'username' in session:
        attendance_records = db.attendance.find()
        return render_template('dashboard.html', attendance_records=attendance_records)
    return redirect(url_for('login'))

@app.route('/register', methods=['GET', 'POST'])
def register():
    if 'username' in session:
        if request.method == 'POST':
            name = request.form['name']
            roll_number = request.form['roll_number']
            image_data = request.form['image']

            # Convert the base64 image data to a numpy array
            image_data = image_data.split(",")[1]
            image_data = np.frombuffer(base64.b64decode(image_data), np.uint8)
            image = cv2.imdecode(image_data, cv2.IMREAD_COLOR)

            # Encode the face
            rgb_image = cv2.cvtColor(image, cv2.COLOR_BGR2RGB)
            face_encodings = face_recognition.face_encodings(rgb_image)
            if face_encodings:
                encoding = face_encodings[0].tolist()

                # Save to MongoDB
                db.face_encodings.insert_one({
                    'name': name,
                    'roll_number': roll_number,
                    'encoding': encoding
                })
                return redirect(url_for('dashboard'))
            else:
                return "No face found. Please try again.", 400

        return render_template('register.html')
    return redirect(url_for('login'))

@app.route('/mark_attendance')
def mark_attendance():
    if 'username' in session:
        return render_template('mark_attendance.html')
    return redirect(url_for('login'))

@app.route('/mark_attendance/facial', methods=['POST'])
def mark_attendance_facial():
    if 'username' in session:
        video_capture = cv2.VideoCapture(0)
        ret, frame = video_capture.read()
        rgb_frame = frame[:, :, ::-1]

        face_locations = face_recognition.face_locations(rgb_frame)
        face_encodings = face_recognition.face_encodings(rgb_frame, face_locations)

        known_face_encodings, known_face_names = get_known_faces()

        for (top, right, bottom, left), face_encoding in zip(face_locations, face_encodings):
            matches = face_recognition.compare_faces(known_face_encodings, face_encoding)
            name = "Unknown"

            face_distances = face_recognition.face_distance(known_face_encodings, face_encoding)
            best_match_index = np.argmin(face_distances)
            if matches[best_match_index]:
                name = known_face_names[best_match_index]
                db.attendance.insert_one({'name': name, 'status': 'present'})

        video_capture.release()
        cv2.destroyAllWindows()
        return redirect(url_for('dashboard'))
    return redirect(url_for('login'))

@app.route('/mark_attendance/speech', methods=['POST'])
def mark_attendance_speech():
    if 'username' in session:
        recognizer = sr.Recognizer()
        microphone = sr.Microphone()

        with microphone as source:
            recognizer.adjust_for_ambient_noise(source)
            audio = recognizer.listen(source)

        try:
            text = recognizer.recognize_google(audio)
            roll_number = text.split()[-1]  # Assumes format "Roll number <num>"

            # Check if the attendee exists in the database
            if db.attendance.find_one({'roll_number': roll_number}):
                return redirect(url_for('dashboard'))  # Attendance already marked
            else:
                db.attendance.insert_one({'roll_number': roll_number, 'status': 'present'})

        except sr.UnknownValueError:
            pass
        except sr.RequestError as e:
            pass

        return redirect(url_for('dashboard'))
    return redirect(url_for('login'))

@app.route('/export')
def export():
    if 'username' in session:
        format = request.args.get('format')
        attendance_records = list(db.attendance.find())
        df = pd.DataFrame(attendance_records)

        # Specify the directory path to save the file
        directory = os.getcwd()  # Get the current working directory

        if format == 'excel':
            # Create the file path
            file_path = os.path.join(directory, "attendance.xlsx")
            # Save the Excel file
            df.to_excel(file_path, index=False)
            return send_file(file_path, as_attachment=True)

        elif format == 'pdf':
            buffer = BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=letter)
            data = [["Name/Roll Number", "Status"]]
            for record in attendance_records:
                name_or_roll_number = record.get('name', record.get('roll_number', 'Unknown'))
                status = record.get('status', 'Unknown')
                data.append([name_or_roll_number, status])

            table = Table(data)
            style = TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ])
            table.setStyle(style)

            elements = [table]
            doc.build(elements)
            buffer.seek(0)
            return send_file(buffer, as_attachment=True, download_name='attendance.pdf', mimetype='application/pdf')

        elif format == 'word':
            document = Document()
            table = document.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Name/Roll Number'
            hdr_cells[1].text = 'Status'

            hdr_cells[0].paragraphs[0].runs[0].bold = True
            hdr_cells[1].paragraphs[0].runs[0].bold = True

            for record in attendance_records:
                row_cells = table.add_row().cells
                row_cells[0].text = str(record.get('name', record.get('roll_number', 'Unknown')))
                row_cells[1].text = str(record.get('status', 'Unknown'))

                for cell in row_cells:
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    cell.paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(12)

            file_path = os.path.join(directory, "attendance.docx")
            document.save(file_path)
            return send_file(file_path, as_attachment=True)

    return redirect(url_for('login'))

@app.route('/logout')
def logout():
    session.pop('username', None)
    return redirect(url_for('index'))

@app.route('/clear_attendance')
def clear_attendance():
    if 'username' in session:
        db.attendance.delete_many({})
        return redirect(url_for('dashboard'))
    return redirect(url_for('login'))
